"""
Generate APA-formatted Word report for the MEPS Healthcare Cost Prediction project.
"""
import os, csv
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

ROOT = Path(__file__).resolve().parent.parent
FIGURES = ROOT / "reports" / "figures"
TABLES = ROOT / "reports" / "tables"
OUTPUT = ROOT / "reports" / "MEPS_Healthcare_Cost_Prediction_Report.docx"

doc = Document()

# ── Global style setup ──────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 2.0

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Times New Roman"
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(14)
    elif level == 2:
        hs.font.size = Pt(13)
    else:
        hs.font.size = Pt(12)

fig_counter = [0]
tbl_counter = [0]


def add_figure(path, caption, width=5.5):
    """Insert a figure with APA-style caption below."""
    if not Path(path).exists():
        p = doc.add_paragraph(f"[Figure not found: {path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    fig_counter[0] += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f"Figure {fig_counter[0]}. ")
    r.italic = True
    r.font.size = Pt(11)
    r2 = cap.add_run(caption)
    r2.italic = True
    r2.font.size = Pt(11)


def add_csv_table(csv_path, title, col_widths=None, fmt_fn=None):
    """Insert a CSV as a formatted Word table with APA caption above."""
    if not Path(csv_path).exists():
        doc.add_paragraph(f"[Table file not found: {csv_path}]")
        return
    tbl_counter[0] += 1
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(4)
    r = cap.add_run(f"Table {tbl_counter[0]}\n")
    r.italic = True
    r.font.size = Pt(11)
    r2 = cap.add_run(title)
    r2.italic = True
    r2.font.size = Pt(11)

    with open(csv_path, newline="") as f:
        reader = list(csv.reader(f))

    headers = reader[0]
    rows = reader[1:]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            if fmt_fn:
                val = fmt_fn(headers[ci], val)
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"

    doc.add_paragraph()


def add_inline_table(headers, rows, title):
    """Insert a manually constructed table."""
    tbl_counter[0] += 1
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(4)
    r = cap.add_run(f"Table {tbl_counter[0]}\n")
    r.italic = True
    r.font.size = Pt(11)
    r2 = cap.add_run(title)
    r2.italic = True
    r2.font.size = Pt(11)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
    doc.add_paragraph()


def fmt_metrics(col, val):
    """Format metric values to 3 decimal places or percentages."""
    try:
        v = float(val)
        if col in ("MAE", "RMSE"):
            return f"${v:,.0f}"
        elif col == "R2":
            return f"{v:.4f}"
        elif col == "MAE_Improvement":
            return f"{v:.1f}%"
        elif col == "N_Features":
            return str(int(v))
        elif "F1" in col or "AUC" in col or "Accuracy" in col:
            return f"{v:.4f}"
        else:
            return f"{v:.4f}"
    except (ValueError, TypeError):
        return val


# ═══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run(
    "Multi-Stage Predictive Modeling of Healthcare Expenditure\n"
    "Using MEPS Longitudinal Data (2013\u20132019)"
)
r.bold = True
r.font.size = Pt(16)
r.font.name = "Times New Roman"

doc.add_paragraph()
for line in [
    "ALY 6980 Experiential Learning",
    "",
    "Northeastern University",
    "College of Professional Studies",
    "",
    "January 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Introduction",
    "2. Research Methods",
    "   2.1 Data Source and Sample",
    "   2.2 Feature Engineering",
    "   2.3 Exploratory Data Analysis",
    "   2.4 Multi-Stage Modeling Architecture",
    "3. Analysis Results and Discussion",
    "   3.1 Data Processing Results",
    "   3.2 Hidden Risers and Warning Signals",
    "   3.3 Patient Phenotype Clustering",
    "   3.4 Stage 1: Risk Tier Classification",
    "   3.5 Stage 1.5: Intermediate Latent Factors",
    "   3.6 Stage 2: Tweedie Expenditure Regression",
    "   3.7 Model Comparison: Impact of Latent Factors",
    "4. Conclusion",
    "5. References",
    "Appendix A: Feature Definitions",
    "Appendix B: DBSCAN Clustering Details",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("1. Introduction", level=1)

doc.add_paragraph(
    "Healthcare expenditure in the United States continues to grow, reaching $4.5 trillion "
    "in 2022 and accounting for approximately 17.3% of GDP (Centers for Medicare & Medicaid "
    "Services [CMS], 2023). Accurately predicting individual-level healthcare costs is "
    "critical for insurance risk adjustment, population health management, and resource "
    "allocation. However, healthcare spending follows a highly skewed distribution: "
    "approximately 5% of the population accounts for over 50% of total expenditures "
    "(Mitchell, 2020), making traditional regression approaches inadequate."
)

doc.add_paragraph(
    "This study addresses these challenges through a multi-stage predictive modeling "
    "framework applied to the Medical Expenditure Panel Survey (MEPS) Household Component "
    "longitudinal data spanning 2013\u20132019. Unlike conventional single-model approaches, "
    "the proposed architecture employs a stepwise supervised learning pipeline where outputs "
    "from earlier classification and latent factor models serve as meta-features for the "
    "final expenditure regression. This design captures nonlinear risk dynamics\u2014such as "
    "hidden cost escalators, mental health trajectories, and healthcare engagement "
    "patterns\u2014that single-stage models typically miss."
)

doc.add_paragraph(
    "The study makes three primary contributions: (1) a three-tier risk classification system "
    "(Stable, Rising, Shock) that replaces the conventional binary high/low dichotomy; "
    "(2) five clinically interpretable intermediate latent factors that bridge raw features "
    "and final cost prediction; and (3) a Tweedie regression model that naturally handles "
    "the zero-inflated, heavy-tailed distribution of healthcare costs, achieving a 32.6% "
    "reduction in Mean Absolute Error compared to the baseline model without latent factors."
)

doc.add_paragraph(
    "Key changes from the previous version (archive_v1) of this project include: expanded "
    "data coverage from Panels 18\u201321 (2013\u20132016) to Panels 18\u201323 (2013\u20132019), "
    "addition of nine medication-derived features from MEPS Prescribed Medicines files, "
    "improved clustering analysis using combined Elbow/Silhouette/DBSCAN methods, and "
    "enhanced model interpretability through SHAP analysis and clinical narrative."
)

# ═══════════════════════════════════════════════════════════════════════
# 2. RESEARCH METHODS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("2. Research Methods", level=1)

# 2.1
doc.add_heading("2.1 Data Source and Sample", level=2)
doc.add_paragraph(
    "Data were obtained from the Agency for Healthcare Research and Quality (AHRQ) Medical "
    "Expenditure Panel Survey (MEPS) Household Component. MEPS employs a nationally "
    "representative overlapping panel design in which each panel is surveyed over two "
    "consecutive calendar years. This study utilized six panels (Panels 18\u201323), "
    "corresponding to the period 2013\u20132019, drawing from three file types per panel:"
)

doc.add_paragraph(
    "\u2022 Longitudinal files: demographics, insurance, self-reported health, chronic "
    "conditions, and annual expenditure totals\n"
    "\u2022 Medical Conditions files: ICD-9/ICD-10 diagnosis codes mapped to Clinical "
    "Classification Software (CCS) categories\n"
    "\u2022 Prescribed Medicines files: medication records including drug class, "
    "quantity, and expenditure"
)

add_inline_table(
    ["Criterion", "Description"],
    [
        ["Age", "Adults aged 18 and older at Year 1"],
        ["Cancer exclusion", "Patients diagnosed with cancer in Year 1 (CANCER_Y1 = 1) excluded"],
        ["Cancer retention", "New Year 2 cancer diagnoses retained (CANCER_Y1 = 0, CANCER_Y2 = 1)"],
        ["Temporal scope", "Year 1 features only; Year 2 used exclusively as outcome"],
    ],
    "Sample Inclusion and Exclusion Criteria",
)

doc.add_paragraph(
    "After applying these criteria, the final analytic sample comprised N = 60,602 "
    "person-panel observations. All cost variables were inflation-adjusted to 2019 dollars "
    "using the Consumer Price Index for Medical Care (CPI-M)."
)

# 2.2
doc.add_heading("2.2 Feature Engineering", level=2)
doc.add_paragraph(
    "A total of 40 features were constructed from Year 1 data, organized into seven "
    "clinically meaningful groups: demographics (age, sex, race/ethnicity, marital "
    "status, poverty category), health status (self-reported physical and mental "
    "health ratings and their longitudinal change), chronic condition flags (diabetes, "
    "hypertension, heart disease, asthma, emphysema, stroke, high cholesterol), "
    "healthcare utilization (office visits, ER visits, inpatient stays, prescriptions), "
    "condition complexity (total conditions count, ill-defined diagnosis count), "
    "medication features (unique drug count, polypharmacy flag, drug class indicators "
    "for cardiovascular, CNS, pain, respiratory, and GI medications, total Rx cost), "
    "and clustering-derived features (K-Means care phenotype, DBSCAN cluster)."
)

doc.add_paragraph(
    'A key innovation is the identification of "Hidden Risers" (Jumpers): patients '
    "whose Year 1 expenditure falls below the 25th percentile yet whose Year 2 "
    "expenditure exceeds the 75th percentile. These patients represent catastrophic "
    "cost escalation that traditional models fail to detect. Feature risk ratios "
    "comparing Jumpers to Stable-Low patients revealed that polypharmacy (risk "
    "ratio = 3.79), GI medications (3.53), and diabetes (3.29) are the strongest "
    "early warning signals for cost escalation."
)

# 2.3
doc.add_heading("2.3 Exploratory Data Analysis", level=2)
doc.add_paragraph(
    "Exploratory analyses examined cost distributions, identified Hidden Risers through "
    "scatter plot analysis (log-log and zoomed linear views), and quantified feature-level "
    "warning signals through relative risk ratios. Unsupervised clustering was performed "
    "using two complementary approaches: K-Means (k = 5, selected via Elbow method with "
    "Silhouette validation and domain constraints) and DBSCAN (eps = 2.619, min_samples = 10), "
    "which identified 6.3% of patients as density-based outliers with mean Year 2 costs "
    "of $29,566\u2014nearly five times the overall average."
)

# 2.4
doc.add_heading("2.4 Multi-Stage Modeling Architecture", level=2)
doc.add_paragraph(
    "The modeling pipeline consists of three sequential stages, where each stage's "
    "outputs serve as inputs to the next:"
)

doc.add_paragraph(
    "Stage 1: Risk Tier Classification. Four classification algorithms (Logistic "
    "Regression, Random Forest, XGBoost, GradientBoosting) were benchmarked to predict "
    "three-tier risk labels (Stable, Rising, Shock). The best model's predicted class "
    "probabilities (PROB_STABLE, PROB_RISING, PROB_SHOCK) and a composite LATENT_RISK_SCORE "
    "were passed forward as meta-features."
)

doc.add_paragraph(
    "Stage 1.5: Intermediate Latent Factor Models. Five clinically motivated sub-models "
    "generate intermediate latent factors:\n"
    "\u2022 1.5A: Mental Health Trajectory (PROB_MH_DECLINE) \u2013 Gradient Boosting classifier "
    "predicting mental health deterioration\n"
    "\u2022 1.5B: Healthcare Engagement Score (ENGAGEMENT_SCORE, PROB_CRISIS_MODE) \u2013 "
    "directly calculated from utilization patterns\n"
    "\u2022 1.5C: Undiagnosed Condition Risk (PROB_UNDIAGNOSED) \u2013 Random Forest predicting "
    "Hidden Riser probability\n"
    "\u2022 1.5D: Cost Escalation Pattern (ESCALATION_SCORE) \u2013 composite of chronic "
    "complexity and utilization indicators\n"
    "\u2022 1.5E: Treatment Discontinuation Hazard (HAZARD_DISCONTINUE) \u2013 Cox Proportional "
    "Hazards survival model"
)

doc.add_paragraph(
    "Stage 2: Tweedie Expenditure Regression. The final model uses Tweedie regression "
    "(compound Poisson-Gamma distribution) to predict raw Year 2 healthcare costs. Tweedie "
    "regression naturally handles the zero mass and heavy right tail of cost distributions "
    "without requiring separate zero/non-zero models. Features were standardized using "
    "StandardScaler, and hyperparameters (power, alpha) were optimized via GridSearchCV."
)

doc.add_paragraph(
    "Data leakage was prevented through four mechanisms: (1) strict temporal feature "
    "separation (Year 1 features only), (2) cancer exclusion to prevent foreknowledge "
    "of catastrophic diagnoses, (3) out-of-fold predictions for all latent factors using "
    "cross-validation, and (4) temporal train/test split (Panels 18\u201322 for training, "
    "Panel 23 for testing)."
)

# ═══════════════════════════════════════════════════════════════════════
# 3. ANALYSIS RESULTS AND DISCUSSION
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("3. Analysis Results and Discussion", level=1)

# 3.1
doc.add_heading("3.1 Data Processing Results", level=2)
doc.add_paragraph(
    "The data processing pipeline successfully merged six MEPS panels into a unified "
    "analytic dataset. Table 1 summarizes the sample construction. The raw combined "
    "dataset contained 93,854 person-panel observations. After restricting to adults "
    "(age \u2265 18), 66,993 records remained. Exclusion of Year 1 cancer patients "
    "(n = 6,391, 9.5%) yielded the final analytic sample of N = 60,602 with 40 features."
)

add_inline_table(
    ["Processing Step", "N"],
    [
        ["Raw combined (Panels 18\u201323)", "93,854"],
        ["After adult filter (age \u2265 18)", "66,993"],
        ["After cancer Y1 exclusion", "60,602"],
    ],
    "Sample Construction Summary",
)

doc.add_paragraph(
    "The target variable, inflation-adjusted Year 2 total expenditure (COST_Y2_ADJ), "
    "exhibited the expected heavy right skew: median = $2,439, mean = $6,843, and "
    "standard deviation = $20,901. Approximately 10% of patients (Shock tier) accounted "
    "for a disproportionate share of total costs. The three-tier distribution was: "
    "Stable (\u226475th percentile, 75.0%), Rising (75th\u201390th, 15.0%), and "
    "Shock (>90th, 10.0%)."
)

# 3.2
doc.add_heading("3.2 Hidden Risers and Warning Signals", level=2)
doc.add_paragraph(
    "Analysis identified 670 Hidden Risers (Jumpers): patients in the bottom quartile "
    "of Year 1 costs who escalated to the top quartile in Year 2. Figure 1 presents "
    "the Year 1 versus Year 2 cost scatter on a log-log scale, revealing the full "
    "range of cost trajectories including extreme outliers. Figure 2 stratifies Jumpers "
    "into three tiers based on Year 2 cost severity."
)

add_figure(
    FIGURES / "hidden_risers_scatter.png",
    "Year 1 vs. Year 2 Healthcare Expenditure (Log-Log Scale). "
    "Red points indicate Hidden Risers (Jumpers). Left panel shows full data range; "
    "right panel zooms to the 99th percentile for detail.",
    width=6.0,
)

add_figure(
    FIGURES / "jumper_tiers_analysis.png",
    "Jumper Tier Analysis by Cost Escalation Magnitude. "
    "Left: scatter plot of Jumpers colored by tier (Moderate, Severe, Extreme). "
    "Right: normalized clinical profile comparison across tiers.",
    width=6.0,
)

doc.add_paragraph(
    "Extreme Jumpers (top 1%, n = 69) had mean Year 2 costs of $165,001 despite "
    "mean Year 1 costs of only $472. Feature risk ratio analysis (Figure 3) revealed "
    "that polypharmacy (RR = 3.79), GI medication use (RR = 3.53), and diabetes "
    "(RR = 3.29) were the strongest Year 1 predictors of Jumper status, suggesting "
    "that medication complexity captures latent disease severity better than "
    "self-reported health ratings."
)

add_figure(
    FIGURES / "exploratory_risk_ratios.png",
    "Feature Warning Signals: Relative Risk Ratios Comparing Jumpers to Stable-Low Patients. "
    "Medication-related features dominate the top positions.",
    width=5.5,
)

# 3.3
doc.add_heading("3.3 Patient Phenotype Clustering", level=2)
doc.add_paragraph(
    "K-Means clustering with k = 5 was selected based on the combined Elbow-Silhouette "
    "analysis (Figure 4). The five clusters represent clinically distinct care phenotypes "
    "ranging from young healthy minimalists (Cluster 0: mean age 37, mean Y2 cost $2,789) "
    "to high-complexity poly-chronic patients (Cluster 4: mean age 61, mean Y2 cost "
    "$33,780). Table 3 provides cluster profiles."
)

add_figure(
    FIGURES / "kmeans_optimization.png",
    "K-Means Cluster Optimization: Inertia (Elbow Method) and Silhouette Score by k. "
    "k = 5 was selected balancing statistical metrics with clinical interpretability.",
    width=5.0,
)

add_csv_table(
    TABLES / "care_cluster_profiles.csv",
    "K-Means Care Phenotype Cluster Profiles (Mean Values)",
)

doc.add_paragraph(
    "DBSCAN clustering (Figure 5) provided a complementary density-based perspective, "
    "identifying 3,825 patients (6.3%) as outliers (Cluster \u22121) with dramatically "
    "elevated costs (mean $29,566). This suggests that a substantial minority of "
    "patients do not conform to typical utilization patterns and may require "
    "individualized risk assessment."
)

add_figure(
    FIGURES / "dbscan_vs_kmeans.png",
    "Comparison of DBSCAN and K-Means Clustering (PCA Projection). "
    "Left: DBSCAN identifies density-based outliers (gray). Right: K-Means assigns "
    "all patients to defined clusters.",
    width=6.0,
)

# 3.4
doc.add_heading("3.4 Stage 1: Risk Tier Classification", level=2)
doc.add_paragraph(
    "Four classification models were benchmarked for three-tier risk prediction using "
    "temporal validation (Panels 18\u201322 for training, Panel 23 for testing). Table 5 "
    "presents the results. Random Forest achieved the highest Macro F1 (0.569), while "
    "XGBoost achieved the highest accuracy (0.762) and ROC-AUC (0.834). Random Forest "
    "was selected as the primary model for latent factor generation due to its superior "
    "balanced performance across all three tiers."
)

add_csv_table(
    TABLES / "stage1_model_benchmark.csv",
    "Stage 1 Model Benchmark: Four-Model Comparison on Test Set (Panel 23)",
    fmt_fn=fmt_metrics,
)

doc.add_paragraph(
    "Figure 6 shows the confusion matrix for the Random Forest model. The Stable "
    "tier was well-classified (recall = 0.75), while Rising (recall = 0.57) and "
    "Shock (recall = 0.46) were more challenging due to class imbalance and the "
    "inherent difficulty of predicting cost escalation from Year 1 features alone."
)

add_figure(
    FIGURES / "confusion_matrix_Random_Forest.png",
    "Confusion Matrix for Random Forest Risk Tier Classification (Stage 1). "
    "Rows represent actual tiers; columns represent predicted tiers.",
    width=4.5,
)

doc.add_paragraph(
    "Feature importance analysis (Figure 7) revealed that Year 1 adjusted cost "
    "(COST_Y1_ADJ, importance = 0.248) was the dominant predictor, followed by "
    "prescription utilization (UTIL_RX_Y1, 0.148) and office visits (UTIL_OB_Y1, "
    "0.112). Notably, the CARE_CLUSTER feature (0.041) and DBSCAN_CLUSTER (0.030) "
    "contributed meaningfully, validating the inclusion of unsupervised clustering "
    "as engineered features."
)

add_figure(
    FIGURES / "feature_importance_stage1.png",
    "Feature Importance from Random Forest Stage 1 Model. "
    "Cost and utilization features dominate, with clustering features providing "
    "additional discriminative power.",
    width=5.5,
)

add_csv_table(
    TABLES / "stage1_feature_importance.csv",
    "Stage 1 Feature Importance Rankings (Random Forest)",
    fmt_fn=fmt_metrics,
)

doc.add_paragraph(
    "The Stage 1 model's predicted class probabilities were combined into a composite "
    "LATENT_RISK_SCORE = PROB_RISING + 2 \u00d7 PROB_SHOCK, providing a continuous risk "
    "measure. Figure 8 demonstrates that this score effectively separates cost "
    "distributions across risk tiers, with Shock-tier patients concentrated at higher "
    "scores and showing substantially elevated Year 2 costs."
)

add_figure(
    FIGURES / "latent_risk_score_analysis.png",
    "Latent Risk Score Analysis (Stage 1 Output). Left: score distribution by actual "
    "risk tier. Right: score vs. actual Year 2 cost, colored by tier.",
    width=5.5,
)

# 3.5
doc.add_heading("3.5 Stage 1.5: Intermediate Latent Factors", level=2)
doc.add_paragraph(
    "Five intermediate latent factors were constructed to capture clinically meaningful "
    "risk dimensions not directly observable in the raw features. These factors bridge "
    "the gap between Stage 1 risk classification and Stage 2 cost prediction."
)

doc.add_heading("3.5.1 Mental Health Trajectory (Stage 1.5A)", level=3)
doc.add_paragraph(
    "A Gradient Boosting classifier predicted mental health decline (PROB_MH_DECLINE) "
    "from Year 1 features. Figure 8 demonstrates the relationship between mental health "
    "risk quintiles and Year 2 cost, with the highest-risk quintile associated with "
    "elevated mean costs."
)
add_figure(
    FIGURES / "stage1_5a_mh_decline.png",
    "Stage 1.5A: Mental Health Decline Risk and Year 2 Cost. "
    "Left: mean cost by risk quintile. Center: cost distribution (low vs. high risk). "
    "Right: summary statistics.",
    width=6.0,
)

doc.add_heading("3.5.2 Healthcare Engagement (Stage 1.5B)", level=3)
doc.add_paragraph(
    "The Healthcare Engagement Score was directly calculated from Year 1 utilization "
    "patterns, distinguishing proactive (regular office visits, medication adherence), "
    "mixed, and crisis-mode (ER-driven) care patterns. Figure 9 shows the engagement "
    "score distribution and its relationship with Year 2 costs."
)
add_figure(
    FIGURES / "stage1_5b_engagement.png",
    "Stage 1.5B: Healthcare Engagement Score Analysis. "
    "Left: engagement vs. cost scatter. Center: mean cost by engagement type. "
    "Right: summary statistics.",
    width=6.0,
)

doc.add_heading("3.5.3 Undiagnosed Condition Risk (Stage 1.5C)", level=3)
doc.add_paragraph(
    "A Random Forest model predicted the probability of being a Hidden Riser "
    "(PROB_UNDIAGNOSED). Figure 10 validates this factor: the Jumper rate increased "
    "across risk quintiles, confirming the model's ability to identify patients at "
    "risk of undetected conditions."
)
add_figure(
    FIGURES / "stage1_5c_undiagnosed.png",
    "Stage 1.5C: Undiagnosed Condition Risk Quintile Analysis. "
    "Left: mean cost by risk quintile. Center: Jumper rate by quintile. "
    "Right: validation metrics.",
    width=6.0,
)

doc.add_heading("3.5.4 Treatment Discontinuation Hazard (Stage 1.5E)", level=3)
doc.add_paragraph(
    "A Cox Proportional Hazards model estimated treatment discontinuation risk "
    "(HAZARD_DISCONTINUE) using medication round data as a proxy for time-to-event. "
    "Figure 11 presents the hazard distribution and Kaplan-Meier curves by risk group."
)
add_figure(
    FIGURES / "stage1_5e_survival_analysis.png",
    "Stage 1.5E: Survival Analysis Results. Clockwise from top-left: "
    "hazard score distribution, hazard vs. Year 2 cost, hazard by risk tier, "
    "and Kaplan-Meier survival curves.",
    width=5.5,
)

doc.add_heading("3.5.5 Latent Factor Intercorrelation", level=3)
doc.add_paragraph(
    "Figure 12 displays the correlation matrix of all latent factors and the Year 2 "
    "cost target. The moderate intercorrelations (r = 0.2\u20130.5) confirm that each "
    "factor captures a distinct risk dimension, supporting their joint inclusion in "
    "the Stage 2 model."
)
add_figure(
    FIGURES / "latent_factors_correlation.png",
    "Correlation Matrix of Intermediate Latent Factors and Year 2 Cost. "
    "Moderate correlations indicate complementary information across factors.",
    width=5.0,
)

# 3.6
doc.add_heading("3.6 Stage 2: Tweedie Expenditure Regression", level=2)
doc.add_paragraph(
    "The Tweedie regression model (power parameter between 1 and 2, representing a "
    "compound Poisson-Gamma distribution) was trained on all 31 features (21 original + "
    "4 Stage 1 latent + 6 Stage 1.5 latent). The Tweedie distribution naturally handles "
    "the zero mass at $0 and the heavy right tail of healthcare cost distributions, "
    "eliminating the need for separate zero/non-zero models."
)

# 3.7
doc.add_heading("3.7 Model Comparison: Impact of Latent Factors", level=2)
doc.add_paragraph(
    "Table 8 presents the central finding of this study: the incremental value of "
    "intermediate latent factors. The baseline Tweedie model using only 21 original "
    "features achieved MAE = $9,114 and near-zero R\u00b2. Adding Stage 1 latent factors "
    "alone provided no improvement, suggesting that simple risk tier probabilities do "
    "not contain sufficient information beyond the raw features. However, adding Stage "
    "1.5 latent factors reduced MAE to $6,140 (a 32.6% improvement) and increased "
    "R\u00b2 to 0.272, demonstrating that clinically motivated intermediate representations "
    "capture risk dynamics invisible to both raw features and simple classification "
    "probabilities."
)

add_csv_table(
    TABLES / "stage2_model_comparison.csv",
    "Stage 2 Model Comparison: Incremental Impact of Latent Factors on Tweedie Regression",
    fmt_fn=fmt_metrics,
)

doc.add_paragraph(
    "This 32.6% MAE reduction represents the primary evidence for the multi-stage "
    "architecture's value. The latent factors that contributed most were the engagement "
    "score (capturing care-seeking behavior), the escalation score (capturing chronic "
    "disease trajectory), and the hazard score (capturing treatment discontinuation "
    "risk). These dimensions are not directly measurable in administrative claims data "
    "but can be inferred through the intermediate modeling stage."
)

# ═══════════════════════════════════════════════════════════════════════
# 4. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("4. Conclusion", level=1)

doc.add_paragraph(
    "This study demonstrates that a multi-stage predictive modeling architecture "
    "substantially outperforms conventional single-stage approaches for healthcare "
    "expenditure prediction. By constructing clinically interpretable intermediate "
    "latent factors\u2014capturing mental health trajectories, healthcare engagement patterns, "
    "undiagnosed condition risk, cost escalation dynamics, and treatment discontinuation "
    "hazard\u2014the model achieved a 32.6% reduction in prediction error compared to a "
    "baseline model using raw features alone."
)

doc.add_paragraph(
    "Key findings include: (1) medication-derived features, particularly polypharmacy "
    "and drug class indicators, are stronger predictors of cost escalation than "
    "self-reported health status; (2) Hidden Risers (low Y1/high Y2 patients) can be "
    "partially identified through Year 1 medication complexity and ill-defined diagnosis "
    "counts; (3) DBSCAN-identified density outliers (6.3% of patients) have mean costs "
    "nearly five times the population average, suggesting that standard risk adjustment "
    "models may systematically underestimate costs for atypical patients; and (4) the "
    "Tweedie distribution provides a natural framework for healthcare cost modeling that "
    "avoids the specification issues of two-part models."
)

doc.add_paragraph(
    "Limitations of this study include: (a) MEPS self-reported conditions may "
    "undercount true disease prevalence; (b) the two-year panel structure limits the "
    "temporal depth available for longitudinal trajectory modeling; (c) the Shock tier "
    "recall of 46% indicates that more than half of true high-cost patients are still "
    "missed at Stage 1, suggesting opportunities for improvement through cost-sensitive "
    "learning or threshold optimization; and (d) medication round data provides a limited "
    "proxy for true time-to-event in the survival analysis component."
)

doc.add_paragraph(
    "Future work should explore: ensemble methods that combine the Tweedie regression "
    "with gradient boosting for cost prediction; richer medication adherence features "
    "from within-year prescription fill patterns; and validation on commercial claims "
    "data to assess generalizability beyond the MEPS survey population."
)

# ═══════════════════════════════════════════════════════════════════════
# 5. REFERENCES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("5. References", level=1)

refs = [
    "Agency for Healthcare Research and Quality. (2023). Medical Expenditure Panel Survey (MEPS) household component. https://meps.ahrq.gov/mepsweb/",
    "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5\u201332. https://doi.org/10.1023/A:1010933404324",
    "Centers for Medicare & Medicaid Services. (2023). National health expenditure data: Historical. https://www.cms.gov/research-statistics-data-and-systems/statistics-trends-and-reports/nationalhealthexpenddata/nationalhealthaccountshistorical",
    "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785\u2013794. https://doi.org/10.1145/2939672.2939785",
    "Dunn, P. K., & Smyth, G. K. (2005). Series evaluation of Tweedie exponential dispersion model densities. Statistics and Computing, 15(4), 267\u2013280. https://doi.org/10.1007/s11222-005-4070-y",
    "Jorgensen, B. (1987). Exponential dispersion models. Journal of the Royal Statistical Society: Series B, 49(2), 127\u2013145.",
    "Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30, 4765\u20134774.",
    "Mitchell, E. M. (2020). Concentration of health expenditures and selected characteristics of high spenders, U.S. civilian noninstitutionalized population, 2017 (Statistical Brief No. 528). Agency for Healthcare Research and Quality.",
    "Mullahy, J. (1998). Much ado about two: Reconsidering retransformation and the two-part model in health econometrics. Journal of Health Economics, 17(3), 247\u2013281. https://doi.org/10.1016/S0167-6296(98)00030-7",
    "Smith, V. A., Maciejewski, M. L., & Olsen, M. K. (2018). Modeling healthcare costs with Tweedie distributions: A comparison of alternative approaches. Health Services and Outcomes Research Methodology, 18(4), 218\u2013234.",
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)

# ═══════════════════════════════════════════════════════════════════════
# APPENDIX A
# ═══════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("Appendix A: Feature Definitions", level=1)

add_inline_table(
    ["Feature", "Description", "Source"],
    [
        ["AGE_Y1", "Age at Year 1 baseline", "Longitudinal"],
        ["SEX", "Biological sex (1=Male, 2=Female)", "Longitudinal"],
        ["RACE", "Race/ethnicity (MEPS RACETHNX)", "Longitudinal"],
        ["MARRY_Y1", "Marital status Year 1", "Longitudinal"],
        ["POVCAT_Y1", "Poverty category Year 1", "Longitudinal"],
        ["INSCOV_Y1", "Insurance coverage Year 1", "Longitudinal"],
        ["RTHLTH_RD1", "Self-rated health Round 1 (1=Excellent to 5=Poor)", "Longitudinal"],
        ["MNHLTH_RD1", "Self-rated mental health Round 1", "Longitudinal"],
        ["MNHLTH_CHANGE", "Mental health change (Round 5 - Round 1)", "Derived"],
        ["COST_Y1_ADJ", "Total expenditure Year 1 (inflation-adjusted)", "Longitudinal"],
        ["UTIL_OB_Y1", "Office-based visits Year 1", "Longitudinal"],
        ["UTIL_ER_Y1", "Emergency room visits Year 1", "Longitudinal"],
        ["UTIL_IP_Y1", "Inpatient stays Year 1", "Longitudinal"],
        ["UTIL_RX_Y1", "Prescription fills Year 1", "Longitudinal"],
        ["CNT_RX_Y1", "Total Rx purchases Year 1", "Prescribed Medicines"],
        ["CNT_UNIQUE_DRUGS", "Count of distinct medications", "Prescribed Medicines"],
        ["POLYPHARMACY_FLAG", "5+ unique drugs (1=Yes)", "Derived"],
        ["TOTAL_RX_COST_Y1", "Total prescription expenditure Year 1", "Prescribed Medicines"],
        ["HAS_CARDIOVASCULAR_RX", "Cardiovascular medication use", "Prescribed Medicines"],
        ["HAS_CNS_RX", "CNS medication use", "Prescribed Medicines"],
        ["HAS_PAIN_RX", "Pain medication use", "Prescribed Medicines"],
        ["HAS_RESPIRATORY_RX", "Respiratory medication use", "Prescribed Medicines"],
        ["HAS_GI_RX", "GI medication use", "Prescribed Medicines"],
        ["CHRONIC_COUNT", "Number of chronic conditions", "Derived"],
        ["CNT_TOTAL_CONDITIONS", "Total medical conditions count", "Medical Conditions"],
        ["CNT_ILL_DEFINED", "Count of ill-defined/symptom diagnoses", "Medical Conditions"],
        ["RATIO_ER_OFFICE", "ER visits / Office visits ratio", "Derived"],
        ["DIAB_Y1_FLAG", "Diabetes diagnosed Year 1", "Longitudinal"],
        ["HIBP_Y1_FLAG", "Hypertension diagnosed Year 1", "Longitudinal"],
        ["CHOL_Y1_FLAG", "High cholesterol diagnosed Year 1", "Longitudinal"],
        ["CARE_CLUSTER", "K-Means care phenotype cluster (0-4)", "Derived"],
        ["DBSCAN_CLUSTER", "DBSCAN density cluster", "Derived"],
    ],
    "Complete Feature Definitions and Data Sources",
)

# ═══════════════════════════════════════════════════════════════════════
# APPENDIX B
# ═══════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("Appendix B: DBSCAN Clustering Details", level=1)

doc.add_paragraph(
    "DBSCAN (Density-Based Spatial Clustering of Applications with Noise) was applied "
    "as a complementary clustering approach to K-Means. Unlike K-Means, which assigns "
    "every point to a cluster, DBSCAN identifies patients who do not belong to any "
    "dense region as noise (Cluster \u22121). These outlier patients often represent "
    "atypical care patterns that warrant individualized assessment."
)

add_figure(
    FIGURES / "dbscan_k_distance.png",
    "DBSCAN Parameter Selection: k-Distance Graph for Epsilon Determination. "
    "The optimal eps = 2.619 was identified at the point of maximum curvature.",
    width=5.0,
)

add_csv_table(
    TABLES / "dbscan_cluster_profiles.csv",
    "DBSCAN Cluster Profiles: Mean Feature Values by Cluster",
)

add_figure(
    FIGURES / "feature_availability.png",
    "Feature Availability Across MEPS Panels. Green indicates the feature is available "
    "in the panel; red indicates it is absent.",
    width=5.5,
)

# ═══════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════
doc.save(str(OUTPUT))
print(f"Report saved to: {OUTPUT}")
print(f"Figures included: {fig_counter[0]}")
print(f"Tables included: {tbl_counter[0]}")
