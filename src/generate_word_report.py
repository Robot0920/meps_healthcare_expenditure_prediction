#!/usr/bin/env python3
"""
Generate Professional Word Report for MEPS Healthcare Cost Prediction Project.

This script creates a formatted Word document with all required sections,
figures, and tables - no code snippets included.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import pandas as pd
from datetime import datetime

# Paths
ROOT = Path(__file__).parent.parent
REPORT_DIR = ROOT / 'reports'
FIGURES_DIR = REPORT_DIR / 'figures'
TABLES_DIR = REPORT_DIR / 'tables'
OUTPUT_PATH = REPORT_DIR / 'MEPS_Healthcare_Cost_Prediction_Report.docx'


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    """Add a heading with consistent formatting."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph_text(doc, text, bold=False, italic=False):
    """Add formatted paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_figure(doc, image_path, caption, fig_num):
    """Add figure with caption."""
    if image_path.exists():
        # Add image
        doc.add_picture(str(image_path), width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add caption
        caption_p = doc.add_paragraph()
        caption_run = caption_p.add_run(f"Figure {fig_num}: {caption}")
        caption_run.italic = True
        caption_run.font.size = Pt(10)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Space after
        return True
    return False


def add_table_from_csv(doc, csv_path, caption, table_num):
    """Add table from CSV file."""
    if csv_path.exists():
        df = pd.read_csv(csv_path)
        
        # Add caption above table
        caption_p = doc.add_paragraph()
        caption_run = caption_p.add_run(f"Table {table_num}: {caption}")
        caption_run.bold = True
        caption_run.font.size = Pt(10)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            header_cells[i].text = str(col)
            set_cell_shading(header_cells[i], '2C3E50')
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = None  # White text
                    run.font.size = Pt(9)
        
        # Data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ''
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        
        doc.add_paragraph()  # Space after
        return True
    return False


def create_report():
    """Generate the complete Word report."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # ========== TITLE PAGE ==========
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("\n\n\n\nMEPS Healthcare Cost Prediction:\nMulti-Stage Risk Modeling Approach")
    title_run.bold = True
    title_run.font.size = Pt(24)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("\n\nIdentifying Hidden Risers Through Step-wise Supervised Learning")
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run(f"\n\n\n\nALY6980 Capstone Project\nNortheastern University\n\n{datetime.now().strftime('%B %Y')}")
    info_run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ========== TABLE OF CONTENTS ==========
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Introduction",
        "2. Research Methods",
        "   2.1 Data Sources",
        "   2.2 Data Processing Pipeline",
        "   2.3 Feature Engineering",
        "   2.4 Modeling Approach",
        "3. Analysis Results and Discussion",
        "   3.1 Exploratory Data Analysis",
        "   3.2 Patient Clustering Analysis",
        "   3.3 Stage 1 Model Performance",
        "   3.4 Feature Importance Analysis",
        "4. Conclusion",
        "5. References",
        "Appendix A: Supplementary Tables",
        "Appendix B: Additional Figures"
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========== 1. INTRODUCTION ==========
    add_heading(doc, "1. Introduction", 1)
    
    intro_text = """
Healthcare expenditure prediction is a critical challenge for insurers, healthcare systems, and policymakers. Traditional actuarial models rely primarily on historical spending patterns, often failing to identify patients who transition from low-cost to high-cost status—the so-called "Hidden Risers." This study addresses this gap by developing a multi-stage predictive modeling framework using the Medical Expenditure Panel Survey (MEPS) longitudinal data.

The primary objective of this research is to predict Year 2 healthcare expenditure using only Year 1 features, with particular emphasis on identifying patients at risk of sudden cost escalation. Unlike binary classification approaches that oversimplify the prediction task, our methodology employs a three-tier risk classification system (Stable, Rising, Shock) that provides more actionable insights for healthcare intervention planning.

This project builds upon and significantly enhances a previous iteration (archive_v1) with the following key improvements:
"""
    doc.add_paragraph(intro_text.strip())
    
    improvements = [
        ("Expanded Data Coverage", "Extended from Panels 18-21 (2013-2016) to Panels 18-23 (2013-2019), providing 60,602 patient records after filtering."),
        ("Enhanced Feature Engineering", "Added 9 medication-derived features from Prescribed Medicines files, including polypharmacy indicators and drug class flags."),
        ("Improved Clustering Analysis", "Upgraded from simple K-Means to a combined approach using Elbow method, Silhouette validation, and DBSCAN for outlier detection."),
        ("Multi-Class Risk Tiers", "Moved from binary classification to a 3-tier system providing more granular risk stratification."),
        ("Interpretability Enhancements", "Added SHAP analysis capability for feature importance explanation and individual prediction interpretation.")
    ]
    
    for title, desc in improvements:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # ========== 2. RESEARCH METHODS ==========
    add_heading(doc, "2. Research Methods", 1)
    
    # 2.1 Data Sources
    add_heading(doc, "2.1 Data Sources", 2)
    data_text = """
This study utilizes data from the Medical Expenditure Panel Survey (MEPS), conducted by the Agency for Healthcare Research and Quality (AHRQ). MEPS is the most comprehensive source of data on healthcare costs and utilization in the United States, featuring a longitudinal panel design that tracks individuals over two consecutive years.

The following MEPS data files were integrated for this analysis:
"""
    doc.add_paragraph(data_text.strip())
    
    # Data files table
    data_table = doc.add_table(rows=1, cols=3)
    data_table.style = 'Table Grid'
    headers = ['File Type', 'File IDs', 'Description']
    header_cells = data_table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        set_cell_shading(header_cells[i], '2C3E50')
        for run in header_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
    
    data_rows = [
        ('Longitudinal', 'h172, h183, h193, h202, h210, h217', 'Panel-specific Y1→Y2 tracking'),
        ('Medical Conditions', 'h162, h170, h180, h190, h199, h207', 'Diagnosis codes (ICD-9/ICD-10)'),
        ('Prescribed Medicines', 'h160a, h168a, h178a, h188a, h197a, h206a', 'Medication records with therapeutic class')
    ]
    for row_data in data_rows:
        row_cells = data_table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            for run in row_cells[i].paragraphs[0].runs:
                run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # 2.2 Data Processing
    add_heading(doc, "2.2 Data Processing Pipeline", 2)
    processing_text = """
The data processing pipeline was designed to create a modeling-ready dataset while preventing data leakage from Year 2 information into Year 1 features. The following steps were implemented:

Cohort Selection Criteria:
"""
    doc.add_paragraph(processing_text.strip())
    
    criteria = [
        "Adults only: AGE_Y1 ≥ 18 years",
        "Cancer exclusion: Patients with CANCER_Y1 = 1 were removed to avoid confounding from known high-cost conditions",
        "Complete data: Required valid Year 1 and Year 2 expenditure records"
    ]
    for c in criteria:
        doc.add_paragraph(c, style='List Bullet')
    
    doc.add_paragraph("This filtering reduced the initial 93,854 records to 60,602 patients in the final cohort.")
    
    # 2.3 Feature Engineering
    add_heading(doc, "2.3 Feature Engineering", 2)
    fe_text = """
A comprehensive set of 40+ features was engineered from the raw MEPS data, organized into the following categories:
"""
    doc.add_paragraph(fe_text.strip())
    
    # Feature categories
    feature_cats = [
        ("Demographics", "Age, sex, race, marital status, poverty category, insurance coverage"),
        ("Health Status", "Self-reported mental health (MNHLTH_RD1), physical health (RTHLTH_RD1), mental health trajectory"),
        ("Chronic Conditions", "Diabetes, hypertension, and cholesterol diagnosis flags; chronic disease count"),
        ("Healthcare Utilization", "ER visits, inpatient stays, office visits, prescription fills"),
        ("Condition Complexity", "Total diagnoses count, ill-defined condition count (ICD R-codes as 'undiagnosed signals')"),
        ("Medication Patterns", "Unique drug count, total Rx cost, polypharmacy flag (≥5 drugs), therapeutic class indicators")
    ]
    
    for cat, desc in feature_cats:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{cat}: ")
        run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    doc.add_paragraph("Inflation Adjustment: All cost variables were adjusted to 2025 dollars using CPI factors to enable cross-year comparisons.")
    
    # 2.4 Modeling Approach
    add_heading(doc, "2.4 Modeling Approach", 2)
    model_text = """
The modeling framework employs a multi-stage supervised learning architecture designed to capture latent risk factors while maintaining clinical interpretability.

Target Variable Definition:
The continuous Year 2 expenditure was transformed into a three-tier risk classification:
"""
    doc.add_paragraph(model_text.strip())
    
    tiers = [
        ("Stable (Tier 0)", "COST_Y2_ADJ ≤ 75th percentile"),
        ("Rising (Tier 1)", "75th < COST_Y2_ADJ ≤ 90th percentile"),
        ("Shock (Tier 2)", "COST_Y2_ADJ > 90th percentile")
    ]
    for tier, defn in tiers:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{tier}: ")
        run.bold = True
        p.add_run(defn)
    
    doc.add_paragraph()
    doc.add_paragraph("Models Benchmarked: Four classification algorithms were evaluated—Logistic Regression, Random Forest, XGBoost, and Gradient Boosting—using temporal validation (earlier panels for training, later panels for testing).")
    
    doc.add_page_break()
    
    # ========== 3. ANALYSIS RESULTS AND DISCUSSION ==========
    add_heading(doc, "3. Analysis Results and Discussion", 1)
    
    # 3.1 EDA
    add_heading(doc, "3.1 Exploratory Data Analysis", 2)
    eda_text = """
Initial exploration revealed significant heterogeneity in healthcare spending patterns. The "Hidden Risers" analysis identified 462 patients (0.8% of cohort) who transitioned from below-median Year 1 costs to above-90th-percentile Year 2 costs—representing the primary target for early intervention.
"""
    doc.add_paragraph(eda_text.strip())
    
    fig_num = 1
    if add_figure(doc, FIGURES_DIR / 'hidden_risers_scatter.png', 
                  "Identifying Hidden Risers: Year 1 vs Year 2 Healthcare Expenditure", fig_num):
        fig_num += 1
    
    doc.add_paragraph("As shown in Figure 1, Hidden Risers (red points) cluster in the lower-left region of Year 1 costs but extend vertically into high Year 2 expenditures, demonstrating that traditional cost-based risk stratification would miss these patients.")
    
    if add_figure(doc, FIGURES_DIR / 'exploratory_risk_ratios.png',
                  "Feature Warning Signals: Relative Risk Ratios for Hidden Risers vs Stable Patients", fig_num):
        fig_num += 1
    
    doc.add_paragraph("Figure 2 presents the relative risk ratios comparing Hidden Risers to Stable Low patients. Key findings include: inpatient utilization (3.2x higher), prescription count (2.9x), chronic disease count (2.2x), and office visits (2.1x). These ratios suggest that utilization intensity, rather than cost alone, serves as an early warning signal.")
    
    # 3.2 Clustering
    add_heading(doc, "3.2 Patient Clustering Analysis", 2)
    cluster_text = """
Unsupervised clustering was employed to identify distinct "Care Phenotypes"—patient segments with similar healthcare utilization patterns. Both K-Means and DBSCAN algorithms were applied and compared.
"""
    doc.add_paragraph(cluster_text.strip())
    
    if add_figure(doc, FIGURES_DIR / 'kmeans_optimization.png',
                  "K-Means Cluster Optimization: Elbow Method with Silhouette Validation", fig_num):
        fig_num += 1
    
    doc.add_paragraph("Figure 3 shows the cluster optimization process. The Elbow method indicated k=5 as the optimal number of clusters, balancing model complexity with interpretability. Silhouette scores remained acceptable (>0.19) across this range.")
    
    if add_figure(doc, FIGURES_DIR / 'dbscan_vs_kmeans.png',
                  "Comparison of DBSCAN and K-Means Clustering Results (PCA Visualization)", fig_num):
        fig_num += 1
    
    doc.add_paragraph("Figure 4 compares DBSCAN and K-Means results. DBSCAN identified 6.7% of patients as noise points (outliers), which may represent atypical high-risk individuals warranting individual review. K-Means assigned all patients to clusters, providing comprehensive segmentation for population health management.")
    
    # Add cluster profile table
    table_num = 1
    if add_table_from_csv(doc, TABLES_DIR / 'care_cluster_profiles.csv',
                         "K-Means Cluster Profiles: Mean Feature Values by Cluster", table_num):
        table_num += 1
    
    doc.add_paragraph("Table 1 presents the cluster profiles, revealing distinct care phenotypes ranging from low-utilization 'Healthy' clusters to high-complexity 'Chronic Disease Management' clusters with elevated costs.")
    
    # 3.3 Model Performance
    add_heading(doc, "3.3 Stage 1 Model Performance", 2)
    perf_text = """
Four machine learning models were trained and evaluated using temporal validation. Performance metrics included Macro F1 (balanced class performance), Weighted F1 (overall accuracy), and ROC-AUC (discrimination ability).
"""
    doc.add_paragraph(perf_text.strip())
    
    if add_table_from_csv(doc, TABLES_DIR / 'stage1_model_benchmark.csv',
                         "Stage 1 Model Benchmark: Performance Comparison Across Algorithms", table_num):
        table_num += 1
    
    doc.add_paragraph("Table 2 shows Random Forest achieved the highest Macro F1 (0.57) and strong balanced performance across all metrics, making it the selected model for Stage 1 risk tier classification.")
    
    if add_figure(doc, FIGURES_DIR / 'confusion_matrix_Random_Forest.png',
                  "Confusion Matrix for Random Forest Risk Tier Classification", fig_num):
        fig_num += 1
    
    doc.add_paragraph("""Figure 5 displays the confusion matrix for the best-performing model. Key observations:
• Stable tier: 75% recall (4886/6531), with some misclassification to Rising tier
• Rising tier: 57% recall (858/1504), the most challenging class to predict
• Shock tier: 46% recall (503/1103), critical for high-cost intervention targeting

The model demonstrates reasonable discrimination while highlighting the inherent difficulty of predicting cost escalation transitions.""")
    
    # 3.4 Feature Importance
    add_heading(doc, "3.4 Feature Importance Analysis", 2)
    fi_text = """
Feature importance analysis reveals which Year 1 characteristics most strongly predict Year 2 risk tier assignment. Both Gini importance and permutation importance methods were applied.
"""
    doc.add_paragraph(fi_text.strip())
    
    if add_figure(doc, FIGURES_DIR / 'feature_importance_stage1.png',
                  "Top 15 Features by Importance for Risk Tier Prediction", fig_num):
        fig_num += 1
    
    if add_table_from_csv(doc, TABLES_DIR / 'stage1_feature_importance.csv',
                         "Feature Importance Scores (Top Features)", table_num):
        table_num += 1
    
    doc.add_paragraph("""Figure 6 and Table 3 reveal the most predictive features:

1. COST_Y1_ADJ (Year 1 expenditure): The strongest predictor, confirming that prior cost remains informative despite our focus on "Hidden Risers"
2. Polypharmacy indicators: Patients with 5+ medications show elevated risk, consistent with clinical literature on adverse drug events
3. CNS/Pain medication use: May indicate mental health burden or chronic pain—known high-cost drivers
4. Office visit utilization: High primary care engagement may reflect disease management complexity

These findings support the clinical validity of our feature engineering approach and suggest actionable intervention targets.""")
    
    doc.add_page_break()
    
    # ========== 4. CONCLUSION ==========
    add_heading(doc, "4. Conclusion", 1)
    conclusion_text = """
This study developed and validated a multi-stage machine learning framework for predicting healthcare cost escalation using MEPS longitudinal data. The key contributions and findings include:

1. Methodological Advancement: The three-tier risk classification (Stable/Rising/Shock) provides more actionable stratification than binary approaches, enabling targeted intervention at different risk levels.

2. Feature Engineering Innovation: Integration of medication-derived features from Prescribed Medicines files—particularly polypharmacy indicators and therapeutic class flags—added clinically meaningful predictors that improved model performance.

3. Hidden Riser Identification: The analysis successfully characterized patients at risk of sudden cost escalation, revealing that utilization intensity (not just current spending) serves as a key early warning signal.

4. Model Performance: Random Forest achieved the best balanced performance (Macro F1: 0.57, Weighted F1: 0.71), demonstrating feasibility of risk tier prediction while acknowledging the inherent difficulty of forecasting healthcare transitions.

5. Interpretability: Cluster analysis identified distinct "Care Phenotypes" that can inform population health management strategies, while feature importance analysis validated the clinical relevance of selected predictors.

Limitations and Future Directions:
• The Rising tier remains challenging to predict, suggesting the need for additional features or alternative modeling approaches
• Integration of SHAP analysis could provide individual-level prediction explanations for clinical decision support
• Stage 2 (attrition/mortality) and Stage 3 (continuous expenditure regression) modeling represent natural extensions of this framework
• Temporal validation across additional years would strengthen generalizability claims

This work demonstrates that combining domain knowledge with modern machine learning techniques can produce interpretable, actionable healthcare cost predictions that go beyond simple actuarial models.
"""
    doc.add_paragraph(conclusion_text.strip())
    
    doc.add_page_break()
    
    # ========== 5. REFERENCES ==========
    add_heading(doc, "5. References", 1)
    references = [
        "Agency for Healthcare Research and Quality. (2024). Medical Expenditure Panel Survey Home. https://meps.ahrq.gov/mepsweb/",
        "Obermeyer, Z., Powers, B., Vogeli, C., & Mullainathan, S. (2019). Dissecting racial bias in an algorithm used to manage the health of populations. Science, 366(6464), 447-453.",
        "Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30.",
        "Cerner Multum. (2024). Multum Lexicon Drug Database. https://www.cerner.com/solutions/drug-database",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785-794.",
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32.",
        "Ester, M., Kriegel, H. P., Sander, J., & Xu, X. (1996). A density-based algorithm for discovering clusters in large spatial databases with noise. KDD-96 Proceedings, 226-231."
    ]
    
    for i, ref in enumerate(references, 1):
        doc.add_paragraph(f"[{i}] {ref}")
    
    doc.add_page_break()
    
    # ========== APPENDIX A ==========
    add_heading(doc, "Appendix A: Supplementary Tables", 1)
    
    if add_table_from_csv(doc, TABLES_DIR / 'dbscan_cluster_profiles.csv',
                         "A1: DBSCAN Cluster Profiles Including Noise Points", table_num):
        table_num += 1
    
    if add_table_from_csv(doc, TABLES_DIR / 'feature_risk_ratios.csv',
                         "A2: Complete Feature Risk Ratios (Hidden Risers vs Stable)", table_num):
        table_num += 1
    
    doc.add_page_break()
    
    # ========== APPENDIX B ==========
    add_heading(doc, "Appendix B: Additional Figures", 1)
    
    if add_figure(doc, FIGURES_DIR / 'feature_availability.png',
                  "B1: Feature Availability Across Processed Dataset", fig_num):
        fig_num += 1
    
    if add_figure(doc, FIGURES_DIR / 'dbscan_k_distance.png',
                  "B2: DBSCAN Epsilon Parameter Selection via K-Distance Graph", fig_num):
        fig_num += 1
    
    if add_figure(doc, FIGURES_DIR / 'latent_risk_score_analysis.png',
                  "B3: Latent Risk Score Distribution by Risk Tier", fig_num):
        fig_num += 1
    
    # Save document
    doc.save(str(OUTPUT_PATH))
    print(f"\n{'='*60}")
    print(f"Report generated successfully!")
    print(f"Output: {OUTPUT_PATH}")
    print(f"{'='*60}")
    
    return OUTPUT_PATH


if __name__ == '__main__':
    create_report()
