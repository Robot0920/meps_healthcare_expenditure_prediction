import sys, os
sys.path.insert(0, '/Users/duanduan/Documents/NEU/ALY6980/healthcare_repo/.pip_pkgs')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
FIGS = '/Users/duanduan/Documents/NEU/ALY6980/healthcare_repo/reports/figures'

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold=False, italic=False, size=12, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_figure(filename, caption, width=5.5):
    path = os.path.join(FIGS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(10)
        run.italic = True
        run.font.name = 'Times New Roman'
        cap.paragraph_format.space_after = Pt(12)
        return True
    else:
        add_para(f'[Figure not found: {filename}]', italic=True, size=10)
        return False


def make_table(headers, data):
    table = doc.add_table(rows=1 + len(data), cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = 'Times New Roman'
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Times New Roman'
    return table


# ================================================================
# [1] TITLE PAGE
# ================================================================
for _ in range(6):
    doc.add_paragraph()

add_para('Multi-Stage Healthcare Cost Prediction', bold=True, size=24,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para('Using Longitudinal MEPS Data (2013\u20132018)', bold=True, size=18,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('Module 7 \u2014 Midterm Report', size=16,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
add_para('ALY 6980 Capstone Project', size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para('Northeastern University', size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para('February 2026', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ================================================================
# [2] EXECUTIVE SUMMARY
# ================================================================
add_heading('Executive Summary', level=1)

add_para(
    'This report presents a multi-stage machine learning pipeline for predicting individual-level '
    'healthcare expenditure using the Medical Expenditure Panel Survey (MEPS) longitudinal data from '
    '2013 to 2018. The study addresses a core challenge in health economics: healthcare costs are '
    'extremely right-skewed, with the top 10% of patients accounting for a disproportionate share of '
    'total spending (Stanton & Rutherford, 2006). Traditional single-model approaches struggle with '
    'this heterogeneity.'
)
add_para(
    'Our pipeline consists of four stages: (1) feature engineering from raw MEPS ASCII files, yielding '
    '31 Year-1 predictors for 60,602 patients; (2) risk tier classification into Stable, Rising, and '
    'Shock categories using Random Forest (ROC-AUC = 0.833); (3) intermediate latent factor extraction '
    'capturing engagement patterns, escalation risk, and medication adherence; and (4) tier-specific '
    'XGBoost regression on log-transformed costs. The tier-specific approach achieved R\u00b2 = 0.600 on '
    'the log scale and R\u00b2 = 0.480 on the dollar scale, substantially outperforming the global '
    'single-model baseline (R\u00b2 = 0.509 log, R\u00b2 = 0.228 dollar). All models use Year-1 data '
    'exclusively, with cross-validated predictions at each stage to prevent data leakage. The ablation '
    'study reveals that the tier-specific modeling strategy is the primary driver of improvement, while '
    'Stage 1.5 latent factors show limited marginal gain, suggesting opportunities for further optimization.'
)

doc.add_page_break()

# ================================================================
# [3] TABLE OF CONTENTS
# ================================================================
add_heading('Table of Contents', level=1)

toc_items = [
    'Executive Summary',
    'Table of Contents',
    '1. Introduction',
    '2. Literature Review',
    '   2.1 Traditional Econometric Approaches',
    '   2.2 Machine Learning Approaches',
    '   2.3 Risk Stratification and Multi-Stage Models',
    '   2.4 Longitudinal Analysis and Feature Engineering',
    '   2.5 Evaluation Metrics for Cost Prediction',
    '3. Research Methods',
    '   3.1 Data Source and Feature Engineering',
    '   3.2 Multi-Stage Pipeline Architecture',
    '   3.3 Data Leakage Prevention',
    '4. Analysis Results and Discussion',
    '   4.1 Stage 1: Risk Tier Classification',
    '   4.2 Stage 1.5: Intermediate Latent Factors',
    '   4.3 Stage 2: Tier-Specific Cost Prediction',
    '   4.4 Ablation Study',
    '5. Conclusion',
    'References',
    'Appendix',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ================================================================
# [4] INTRODUCTION
# ================================================================
add_heading('1. Introduction', level=1)

add_para(
    'Healthcare expenditure prediction is a fundamental problem in health economics and actuarial '
    'science. In the United States, healthcare spending reached $4.5 trillion in 2022, representing '
    '17.3% of GDP (Centers for Medicare & Medicaid Services [CMS], 2023). The distribution of this '
    'spending is highly concentrated: approximately 5% of the population accounts for over 50% of '
    'total expenditures (Stanton & Rutherford, 2006; Mitchell & Machlin, 2014). Accurate prediction '
    'of individual-level costs enables insurers, healthcare systems, and policymakers to allocate '
    'resources efficiently, design targeted interventions, and manage financial risk.'
)
add_para(
    'However, healthcare cost prediction presents unique statistical challenges. The distribution of '
    'costs is severely right-skewed, with a large mass at zero and extreme right-tail values exceeding '
    '$500,000 annually (Manning & Mullahy, 2001). Traditional linear models and even standard machine '
    'learning approaches often fail to capture this heterogeneity. Two-part models (Duan et al., 1983), '
    'generalized linear models with Tweedie distributions (J\u00f8rgensen, 1987), and more recently, '
    'tree-based ensemble methods (Bertsimas et al., 2008) have been proposed, each with trade-offs '
    'between interpretability and predictive accuracy.'
)
add_para(
    'This study leverages the Medical Expenditure Panel Survey (MEPS), a nationally representative '
    'longitudinal survey conducted by the Agency for Healthcare Research and Quality (AHRQ). MEPS '
    'follows individuals over two-year overlapping panels, collecting detailed information on '
    'demographics, health status, chronic conditions, healthcare utilization, and expenditures '
    '(Cohen et al., 2009). We use six panels (18\u201323) covering 2013\u20132018, yielding 60,602 '
    'patient records with 31 Year-1 features.'
)
add_para(
    'Our contribution is a multi-stage pipeline that decomposes the prediction problem into sequential '
    'sub-tasks: (1) risk tier classification, (2) latent clinical factor extraction, and (3) tier-specific '
    'cost regression. This approach is motivated by the observation that a single global model spends most '
    'of its capacity learning the gross differences between low-cost and high-cost patients\u2014a distinction '
    'that can be made more efficiently through classification\u2014while tier-specific models can focus on '
    'the more nuanced within-tier cost variation. All features are restricted to Year-1 data, and '
    'cross-validated predictions are used at each intermediate stage to prevent data leakage '
    '(Kaufman et al., 2012).'
)

doc.add_page_break()

# ================================================================
# [5] LITERATURE REVIEW
# ================================================================
add_heading('2. Literature Review', level=1)

add_heading('2.1 Traditional Econometric Approaches', level=2)
add_para(
    'The foundational work of Duan et al. (1983) introduced the two-part model for healthcare costs: '
    'a logistic regression for the probability of any expenditure, followed by a log-linear regression '
    'for the conditional cost given positive spending. Manning and Mullahy (2001) extended this framework '
    'by comparing log-scale OLS with generalized linear models (GLMs), finding that the choice of '
    'retransformation method significantly affects prediction accuracy. J\u00f8rgensen (1987) introduced '
    'the Tweedie distribution family, which naturally handles the zero-inflated continuous distribution '
    'of healthcare costs through a compound Poisson-Gamma formulation.'
)

add_heading('2.2 Machine Learning Approaches', level=2)
add_para(
    'Bertsimas et al. (2008) demonstrated that classification and regression trees (CART) and ensemble '
    'methods could outperform traditional econometric models for healthcare cost prediction, particularly '
    'for capturing nonlinear relationships and interactions. Rose (2016) compared multiple machine learning '
    'algorithms on MEPS data and found that Super Learner ensembles achieved superior performance. More '
    'recently, gradient boosting methods such as XGBoost (Chen & Guestrin, 2016) and LightGBM '
    '(Ke et al., 2017) have become the state of the art for tabular prediction tasks, including healthcare '
    'cost modeling.'
)
add_para(
    'Morid et al. (2017) provided a comprehensive review of supervised machine learning methods for '
    'predicting healthcare costs, finding that ensemble methods consistently outperformed individual '
    'algorithms. Tamang et al. (2017) demonstrated the value of incorporating medication-level features '
    'for cost prediction, a finding that informs our feature engineering approach.'
)

add_heading('2.3 Risk Stratification and Multi-Stage Models', level=2)
add_para(
    'The concept of patient risk stratification has been widely adopted in population health management. '
    'The Johns Hopkins ACG system (Weiner et al., 1991) and the CMS-HCC model (Pope et al., 2004) use '
    'diagnosis-based risk scores to predict future costs. Ash et al. (2000) introduced the DxCG model '
    'using hierarchical condition categories for risk adjustment. Our approach differs by learning risk '
    'strata directly from the data rather than using pre-defined clinical groupings, and by explicitly '
    'modeling the transition between stages.'
)

add_heading('2.4 Longitudinal Analysis and Feature Engineering', level=2)
add_para(
    'MEPS longitudinal structure has been exploited in several studies. Zuvekas and Olin (2009) examined '
    'persistence in healthcare expenditures across years, finding substantial autocorrelation (r \u2248 0.5 '
    'between consecutive years). This motivates the inclusion of Year-1 cost as a predictor. Medication '
    'adherence\u2014measured through prescription fill patterns\u2014has been shown to predict cost escalation '
    'in chronic disease populations (Roebuck et al., 2011; Sokol et al., 2005). Our Stage 1.5 medication '
    'adherence model draws on survival analysis concepts to quantify discontinuation risk from Year-1 '
    'prescription timelines.'
)

add_heading('2.5 Evaluation Metrics for Cost Prediction', level=2)
add_para(
    'The choice of evaluation metric is critical for healthcare cost models. R\u00b2 on raw dollar amounts '
    'is dominated by extreme outliers, leading Bertsimas et al. (2008) and Rose (2016) to recommend '
    'evaluation on the log-transformed scale. Mean absolute error (MAE) provides a more interpretable '
    'measure of average prediction quality. Basu and Manning (2009) argued for evaluating models on '
    'multiple scales. Following this guidance, we report R\u00b2 on both log and dollar scales, along with MAE.'
)

doc.add_page_break()

# ================================================================
# [6] RESEARCH METHODS
# ================================================================
add_heading('3. Research Methods', level=1)

add_heading('3.1 Data Source and Feature Engineering', level=2)
add_para(
    'The data source is the MEPS Household Component, panels 18\u201323, spanning 2013\u20132018. Raw data '
    'was obtained as fixed-width ASCII files from the AHRQ MEPS data repository and parsed into structured '
    'Parquet format. The final analytic dataset contains 60,602 patient records with 39 variables, of which '
    '31 are used as predictive features.'
)
add_para('Feature categories include:', bold=True)
add_para(
    '\u2022 Demographics (6): age, sex, race, marital status, poverty category, insurance coverage\n'
    '\u2022 Health Status (2): self-rated mental health (MNHLTH_RD1), self-rated physical health (RTHLTH_RD1)\n'
    '\u2022 Chronic Conditions (4): diabetes, hypertension, high cholesterol flags, comorbidity count\n'
    '\u2022 Utilization (4): ER visits, inpatient stays, office visits, prescription fills\n'
    '\u2022 Engineered Features (3): total conditions, ill-defined symptom count, ER-to-office ratio\n'
    '\u2022 Medication (9): Rx count, unique drugs, Rx cost, polypharmacy flag, therapeutic class flags\n'
    '\u2022 Clusters (2): K-Means and DBSCAN patient clusters from utilization patterns\n'
    '\u2022 Cost (1): inflation-adjusted Year-1 total expenditure (COST_Y1_ADJ)'
)
add_para(
    'All costs are inflation-adjusted to 2025 dollars using CPI medical care factors. The target variable '
    'is COST_Y2_ADJ (Year-2 inflation-adjusted expenditure). An intermediate target, RISK_TIER, is defined '
    'based on COST_Y2_ADJ percentiles: Stable (\u226475th percentile, ~75%), Rising (75th\u201390th, ~15%), '
    'and Shock (>90th, ~10%).'
)

add_heading('3.2 Multi-Stage Pipeline Architecture', level=2)
add_para(
    'Stage 1 (Risk Tier Classification): Predicts RISK_TIER from Year-1 features using Random Forest. '
    'Outputs tier probabilities (PROB_STABLE, PROB_RISING, PROB_SHOCK) and a continuous '
    'LATENT_RISK_SCORE = 0\u00d7P(Stable) + 1\u00d7P(Rising) + 2\u00d7P(Shock).'
)
add_para(
    'Stage 1.5 (Latent Factor Extraction): Constructs four intermediate clinical variables\u2014'
    'ENGAGEMENT_SCORE, PROB_UNDIAGNOSED, ESCALATION_SCORE, and HAZARD_DISCONTINUE\u2014each capturing '
    'a specific dimension of risk not directly expressed by the original features.'
)
add_para(
    'Stage 2 (Tier-Specific Cost Prediction): Trains separate XGBoost regressors for each risk tier on '
    'log(1 + COST_Y2_ADJ), using original features, interaction terms, Stage 1 risk score, and Stage 1.5 '
    'latent factors as inputs.'
)

add_heading('3.3 Data Leakage Prevention', level=2)
add_para(
    'Strict protocols prevent information leakage (Kaufman et al., 2012): (a) all features use Year-1 '
    'data exclusively; (b) Stage 1 tier probabilities are generated via out-of-fold prediction; '
    '(c) Stage 1.5 latent factors use 5-fold cross_val_predict; (d) the train/validation/test split '
    '(70/20/10) is stratified by MEPS panel to maintain temporal representativeness.'
)

doc.add_page_break()

# ================================================================
# [7] ANALYSIS RESULTS AND DISCUSSION
# ================================================================
add_heading('4. Analysis Results and Discussion', level=1)

# --- 4.1 Stage 1 ---
add_heading('4.1 Stage 1: Risk Tier Classification', level=2)

add_para('Four classifiers were benchmarked on the risk tier classification task. Table 1 presents the results.')

make_table(
    ['Model', 'Accuracy', 'F1 (Macro)', 'ROC-AUC'],
    [
        ['Random Forest', '0.695', '0.579', '0.833'],
        ['XGBoost', '0.768', '0.553', '0.836'],
        ['GradientBoosting', '0.762', '0.542', '0.833'],
        ['Logistic Regression', '0.684', '0.550', '0.815'],
    ]
)
add_para('Table 1. Stage 1 Model Benchmark Results (Test Set, n = 9,138)', italic=True, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(
    'Random Forest was selected based on the highest F1 Macro score (0.579), which weights all three '
    'classes equally. This is appropriate because the Shock class (10%) is the primary intervention '
    'target. The raw accuracy of 69.5% is below the naive majority-class baseline of 75%; this '
    'trade-off is intentional, sacrificing majority-class precision for improved minority-class '
    'detection (He & Garcia, 2009). Per-class analysis reveals: Stable (precision = 0.90, recall = 0.76), '
    'Rising (precision = 0.33, recall = 0.59), and Shock (precision = 0.56, recall = 0.44). The Shock '
    'recall of 44% indicates that more than half of true high-cost patients are missed.'
)

add_figure('confusion_matrix_Random_Forest.png',
           'Figure 1. Confusion Matrix \u2014 Random Forest Risk Tier Classification', 4.0)

add_para(
    'SHAP analysis (Lundberg & Lee, 2017) identified the top predictive features for Shock classification. '
    'Prior-year cost (COST_Y1_ADJ, mean |SHAP| = 0.061) dominates, consistent with the cost persistence '
    'effect (Zuvekas & Olin, 2009). Prescription volume (UTIL_RX_Y1 = 0.032), office visits '
    '(UTIL_OB_Y1 = 0.023), and medication cost (TOTAL_RX_COST_Y1 = 0.021) follow, reflecting chronic '
    'disease treatment intensity.'
)

add_figure('shap_summary_shock.png',
           'Figure 2. SHAP Feature Importance for Shock (Tier 2) Prediction', 5.0)

add_para(
    'The LATENT_RISK_SCORE (mean = 0.712, SD = 0.420) was validated by examining mean Year-2 cost '
    'across score deciles (Figure 3). The monotonically increasing pattern confirms correct rank-ordering.'
)

add_figure('latent_risk_score_analysis.png',
           'Figure 3. Latent Risk Score Validation: Distribution, Decile Means, and Hexbin', 5.5)

doc.add_page_break()

# --- 4.2 Stage 1.5 ---
add_heading('4.2 Stage 1.5: Intermediate Latent Factors', level=2)

add_para(
    'Stage 1.5 constructs four intermediate variables, each capturing a clinical dimension not directly '
    'expressed by the original features. A fifth factor (PROB_MH_DECLINE) was removed after evaluation '
    'showed negligible correlation with cost (r = 0.01).'
)

make_table(
    ['Latent Factor', 'Method', 'Key Inputs', 'r with Cost'],
    [
        ['ENGAGEMENT_SCORE', 'Direct formula', 'ER, Office, IP ratios', '0.235'],
        ['PROB_UNDIAGNOSED', 'RF + 5-fold CV', 'Ill-defined Dx, utilization', '\u22120.207'],
        ['ESCALATION_SCORE', 'Composite index', 'Chronic count, IP, polypharmacy', '0.289'],
        ['HAZARD_DISCONTINUE', 'GB + 5-fold CV', 'Rx fill density, gap ratio', '0.384'],
    ]
)
add_para('Table 2. Stage 1.5 Latent Factors Summary', italic=True, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(
    'The HAZARD_DISCONTINUE variable, inspired by survival analysis methodology (Cox, 1972; '
    'Roebuck et al., 2011), shows the strongest correlation with Year-2 cost among Stage 1.5 factors '
    '(r = 0.384). It quantifies medication adherence risk by analyzing Year-1 prescription fill '
    'timelines: fill density (proportion of months with fills), gap ratio (fraction of time between '
    'fills that are gaps), and trailing gap (months since last fill at year-end). Figure 4 demonstrates '
    'that patients in the worst adherence quintile (Q5) have a median Year-2 cost of $9,790 compared '
    'to near-zero for the best adherence quintile (Q1)\u2014a ratio of approximately 9,790:1.'
)

add_figure('stage1_5e_survival_analysis.png',
           'Figure 4. Medication Adherence Risk Analysis (Stage 1.5E)', 5.5)

add_para(
    'The correlation matrix (Figure 5) reveals important structural relationships. PROB_CRISIS_MODE '
    'and ENGAGEMENT_SCORE are almost perfectly inversely correlated (r = \u22120.99), indicating redundancy. '
    'HAZARD_DISCONTINUE correlates strongly with LATENT_RISK_SCORE (r = 0.91), suggesting substantial '
    'overlap with Stage 1 outputs. These findings indicate opportunities for factor refinement.'
)

add_figure('latent_factors_correlation.png',
           'Figure 5. Correlation Matrix: All Latent Factors and Target Variable', 4.5)

doc.add_page_break()

# --- 4.3 Stage 2 ---
add_heading('4.3 Stage 2: Tier-Specific Cost Prediction', level=2)

add_para(
    'Stage 2 trains XGBoost regressors on log(1 + COST_Y2_ADJ) using 33 features: 21 original, '
    '6 interaction terms, LATENT_RISK_SCORE, and 5 Stage 1.5 latent factors. Two approaches are compared.'
)

make_table(
    ['Model', 'R\u00b2 (log)', 'R\u00b2 (dollar)', 'MAE'],
    [
        ['Global XGBoost', '0.509', '0.228', '$5,046'],
        ['Tier-Specific XGBoost', '0.600', '0.480', '$3,440'],
    ]
)
add_para('Table 3. Stage 2 Overall Model Comparison (Test Set, n = 6,066)', italic=True, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(
    'The tier-specific approach achieves R\u00b2 = 0.600 on the log scale, a substantial improvement '
    'over the global baseline (R\u00b2 = 0.509). The dollar-scale R\u00b2 improves from 0.228 to 0.480, '
    'and MAE decreases from $5,046 to $3,440.'
)

make_table(
    ['Tier', 'n (test)', 'R\u00b2 (log)', 'R\u00b2 (dollar)', 'MAE'],
    [
        ['Stable', '4,552', '0.372', '0.095', '$808'],
        ['Rising', '906', '0.046', '0.023', '$2,719'],
        ['Shock', '608', '0.136', '0.030', '$24,213'],
    ]
)
add_para('Table 4. Within-Tier Performance of Tier-Specific XGBoost', italic=True, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(
    'Within-tier analysis reveals heterogeneous performance. The Stable tier achieves R\u00b2 = 0.372 '
    '(log) with MAE of only $808. The Shock tier R\u00b2 = 0.136 (log) with MAE = $24,213 operates '
    'on a cost range of $10K\u2013$900K. The Rising tier shows the weakest performance (R\u00b2 = 0.046), '
    'consistent with its nature as a transitional category.'
)

add_figure('stage2_results.png',
           'Figure 6. Stage 2 Tier-Specific Model Results (6-Panel Summary)', 5.8)

doc.add_page_break()

# --- 4.4 Ablation ---
add_heading('4.4 Ablation Study', level=2)

make_table(
    ['Feature Set', 'R\u00b2 (log)', 'R\u00b2 (dollar)', 'MAE'],
    [
        ['Original Only (21)', '0.601', '0.480', '$3,444'],
        ['+ Interactions (27)', '0.600', '0.481', '$3,441'],
        ['+ LATENT_RISK_SCORE (28)', '0.600', '0.475', '$3,447'],
        ['+ Stage 1.5 Full (33)', '0.600', '0.480', '$3,433'],
    ]
)
add_para('Table 5. Ablation Study: Feature Set Impact on Tier-Specific Model Performance',
         italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(
    'A critical finding emerges: within the tier-specific framework, adding Stage 1.5 latent factors '
    'provides minimal marginal improvement over original features alone (R\u00b2 log: 0.601 \u2192 0.600). '
    'The primary performance gain comes from the tier-specific modeling strategy itself (Global R\u00b2 = 0.509 '
    'vs. Tier-Specific R\u00b2 = 0.600), not from the intermediate latent variables. This suggests that '
    'once patients are stratified by tier, the original Year-1 features\u2014particularly COST_Y1_ADJ\u2014'
    'already capture most of the predictable within-tier variation. The Stage 1.5 factors, while clinically '
    'interpretable, are either redundant with the tier-stratification signal (HAZARD_DISCONTINUE, r = 0.91 '
    'with LATENT_RISK_SCORE) or too weak to add incremental value (PROB_UNDIAGNOSED, r = \u22120.207).'
)

add_figure('stage2_ablation.png',
           'Figure 7. Ablation Study: R\u00b2 and MAE by Feature Configuration', 5.5)

add_para(
    'This finding has important implications for pipeline refinement. Rather than wrapping clinical '
    'signals into intermediate classifiers, future work should consider feeding the raw sub-features '
    '(e.g., Rx fill density, gap ratio, trailing gap) directly into Stage 2 models, allowing XGBoost '
    'to learn the optimal nonlinear transformations itself.'
)

doc.add_page_break()

# ================================================================
# [8] CONCLUSION
# ================================================================
add_heading('5. Conclusion', level=1)

add_para(
    'This study demonstrates that a multi-stage pipeline\u2014specifically the tier-specific modeling '
    'strategy\u2014substantially improves healthcare cost prediction compared to a single global model. '
    'The tier-specific XGBoost achieves R\u00b2 = 0.600 (log scale) and R\u00b2 = 0.480 (dollar scale) '
    'on individual-level MEPS data, compared to 0.509 and 0.228 for the global baseline. The ablation '
    'study reveals that the tier-stratification architecture, rather than the Stage 1.5 latent factors, '
    'is the primary driver of improvement. Key limitations include low Shock recall (44%) in Stage 1, '
    'weak Rising-tier prediction (within-tier R\u00b2 = 0.046), and redundancy among latent factors. '
    'Future work will focus on pruning redundant variables, incorporating raw sub-features directly, '
    'Bayesian hyperparameter optimization, and external validation on held-out panels.'
)

doc.add_page_break()

# ================================================================
# [9] REFERENCES
# ================================================================
add_heading('References', level=1)

refs = [
    'Ash, A. S., Ellis, R. P., Pope, G. C., Ayanian, J. Z., Bates, D. W., Burstin, H., ... & Yu, W. (2000). Using diagnoses to describe populations and predict costs. Health Care Financing Review, 21(3), 7\u201328.',
    'Basu, A., & Manning, W. G. (2009). Issues for the next generation of health care cost analyses. Medical Care, 47(7 Suppl 1), S109\u2013S114.',
    'Bertsimas, D., Bjarnad\u00f3ttir, M. V., Kane, M. A., Kryder, J. C., Pandey, R., Vber, S., & Wang, G. (2008). Algorithmic prediction of health-care costs. Operations Research, 56(6), 1382\u20131392.',
    'Centers for Medicare & Medicaid Services. (2023). National Health Expenditure Data. CMS.gov.',
    'Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785\u2013794.',
    'Cohen, S. B., DiGaetano, R., & Goksel, H. (2009). Estimation procedures in the 1996 Medical Expenditure Panel Survey Household Component. AHRQ Methodology Report No. 2009-03.',
    'Cox, D. R. (1972). Regression models and life-tables. Journal of the Royal Statistical Society: Series B, 34(2), 187\u2013202.',
    'Duan, N., Manning, W. G., Morris, C. N., & Newhouse, J. P. (1983). A comparison of alternative models for the demand for medical care. Journal of Business & Economic Statistics, 1(2), 115\u2013126.',
    'He, H., & Garcia, E. A. (2009). Learning from imbalanced data. IEEE Transactions on Knowledge and Data Engineering, 21(9), 1263\u20131284.',
    'J\u00f8rgensen, B. (1987). Exponential dispersion models. Journal of the Royal Statistical Society: Series B, 49(2), 127\u2013162.',
    'Kaufman, S., Rosset, S., Perlich, C., & Stitelman, O. (2012). Leakage in data mining: Formulation, detection, and avoidance. ACM Transactions on Knowledge Discovery from Data, 6(4), 1\u201321.',
    'Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., ... & Liu, T. Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems, 30, 3146\u20133154.',
    'Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30, 4765\u20134774.',
    'Manning, W. G., & Mullahy, J. (2001). Estimating log models: To transform or not to transform? Journal of Health Economics, 20(4), 461\u2013494.',
    'Mitchell, E. M., & Machlin, S. R. (2014). Concentration of health expenditures and selected characteristics of high spenders, U.S. civilian noninstitutionalized population, 2014. AHRQ Statistical Brief #497.',
    'Morid, M. A., Kawamoto, K., Ault, T., Dorius, J., & Abdelrahman, S. (2017). Supervised learning methods for predicting healthcare costs: Systematic literature review and empirical evaluation. AMIA Annual Symposium Proceedings, 1312\u20131321.',
    'Pope, G. C., Kautter, J., Ellis, R. P., Ash, A. S., Ayanian, J. Z., Iezzoni, L. I., ... & Robst, J. (2004). Risk adjustment of Medicare capitation payments using the CMS-HCC model. Health Care Financing Review, 25(4), 119\u2013141.',
    'Roebuck, M. C., Liberman, J. N., Gemmill-Toyama, M., & Brennan, T. A. (2011). Medication adherence leads to lower health care use and costs despite increased drug spending. Health Affairs, 30(1), 91\u201399.',
    'Rose, S. (2016). A machine learning framework for plan payment risk adjustment. Health Services Research, 51(6), 2358\u20132374.',
    'Sokol, M. C., McGuigan, K. A., Verbrugge, R. R., & Epstein, R. S. (2005). Impact of medication adherence on hospitalization risk and healthcare cost. Medical Care, 43(6), 521\u2013530.',
    'Stanton, M. W., & Rutherford, M. K. (2006). The high concentration of U.S. health care expenditures. AHRQ Research in Action, Issue 19.',
    'Tamang, S., Milstein, A., S\u00f8rensen, H. T., Pedersen, L., Mackey, L., Betterton, J. R., ... & Hernandez-Boussard, T. (2017). Predicting patient "cost blooms" in Denmark: A longitudinal population-based study. BMJ Open, 7(1), e011580.',
    'Weiner, J. P., Starfield, B. H., Steinwachs, D. M., & Mumford, L. M. (1991). Development and application of a population-oriented measure of ambulatory care case-mix. Medical Care, 29(5), 452\u2013472.',
    'Zuvekas, S. H., & Olin, G. L. (2009). Validating household reports of health care use in the Medical Expenditure Panel Survey. Health Services Research, 44(5 Pt 1), 1679\u20131700.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ================================================================
# [10] APPENDIX
# ================================================================
add_heading('Appendix', level=1)

add_heading('A. Feature Importance \u2014 Stage 1', level=2)
add_figure('feature_importance_stage1.png',
           'Figure A1. Gini Feature Importance \u2014 Random Forest Stage 1 Model', 5.0)

add_heading('B. Clustering Analysis', level=2)
add_figure('kmeans_optimization.png',
           'Figure A2. K-Means Elbow Method for Optimal Cluster Selection', 4.5)
add_figure('dbscan_vs_kmeans.png',
           'Figure A3. DBSCAN vs K-Means Cluster Comparison', 5.5)

add_heading('C. Hidden Riser Analysis', level=2)
add_figure('hidden_risers_scatter.png',
           'Figure A4. Hidden Riser (Jumper) Scatter Analysis \u2014 Year-1 vs Year-2 Cost', 5.5)
add_figure('jumper_tiers_analysis.png',
           'Figure A5. Jumper Tier Distribution Analysis', 5.5)

add_heading('D. Exploratory Risk Ratios', level=2)
add_figure('exploratory_risk_ratios.png',
           'Figure A6. Feature-Level Risk Ratios for Shock Tier', 5.0)

output = '/Users/duanduan/Documents/NEU/ALY6980/healthcare_repo/reports/Module_7_Midterm_Report.docx'
doc.save(output)
print(f'Saved: {output}')
print(f'References: {len(refs)}')
